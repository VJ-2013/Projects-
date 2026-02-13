import os 
import traceback
from pathlib import Path 
from PyPDF2 import PdfReader
from openpyxl import load_workbook
from docx import Document

extract_partial=True  # True = fetch only some lines, False = fetch full content 
max_lines=5  # Limit for partial extraction

Base_Dir=Path(__file__).resolve().parent
Source_Dir=Base_Dir/"Source Files"
output_path=Base_Dir/"output.docx"

def readTXT(filepath:Path) -> str:
    with open(filepath,"r") as TXT:
        ContentTXT = TXT.readlines()
        if extract_partial:
            ContentTXT=ContentTXT[:max_lines]
        return"".join(ContentTXT).strip()
        # return "".join(ContentTXT[:max_lines]) if extract_partial else "".join(ContentTXT)......
    
def readDocx(filepath:Path) -> str:
    DOCX = Document(filepath)
    ContentDOCX = [p.text for p in DOCX.paragraphs if p.text.strip()
    return "\n".join(ContentDOCX[:max_lines]) if extract_partial else"\n".join(ContentDOCX)

def readPDF(filepath:Path) -> str:
    PDF = PdfReader(filepath)
    Pages = PDF.pages
    Limit = max_lines if extract_partial else len(Pages)
    ContentPDF = [Pages[i].extract_text() or "\n" for i in range(min(Limit,len(Pages)))]
    return "\n".join(ContentPDF)
   
def readExcel(filepath:Path) -> str:
    XL = load_workbook(filepath)
    ContentXL=[]
    for i in XL.sheetnames:
        insideXL=XL[i]
        ContentXL.append(i)
        for i, row in enumerate(insideXL.iter_rows(values_only=True), start=1):
            if extract_partial and i>max_lines:
                break
            ContentXL.append("    |    ".join("" if cell is None else str (cell) for cell in row ))
    return "\n".join(ContentXL)    

def Merge_Files(inputfolder:Path, outputpath:Path):
    print("Merge Start.......✅")
    print(f"Script directory : {Base_Dir}")
    print(f"Working directory: {Path.cwd()}")
    print(f"Source folder    : {inputfolder}")
    print(f"Output path      : {outputpath}")
    print("...................")
    DOC = Document()
    DOC.add_heading("💵💰MERGED VAANI'S BANK DATA💲", level=0)
    Supported_exts = {".txt",".docx",".pdf",".xlsx"}
    files = sorted([i for i in inputfolder.iterdir() if i.is_file() and i.suffix.lower() in Supported_exts], key=lambda i: i.name.lower(),) 
    if not files:
        print("No supported FILES FOUND IN source files folder!!......")
    else:
        for filepath in files:
            print(f" -> Reading: {filepath.name}")
            DOC.add_heading(filepath.name, level=1)
            try:
                CALLINGMETHODS = filepath.suffix.lower()
                if CALLINGMETHODS == ".txt" :
                    Content= readTXT(filepath)
                elif CALLINGMETHODS == ".docx" :
                    Content= readDocx(filepath)
                elif CALLINGMETHODS == ".pdf" :
                    Content= readPDF(filepath)
                elif CALLINGMETHODS == ".xlsx" :
                    Content= readExcel(filepath)
                else:
                    Content = "Unsupported file type............."
            except:
                print(f"Error reading {filepath.name}")
                Content="Error reading file!!✖️"
            if not Content.strip():
                Content="No readable content!!🚫"
            DOC.add_paragraph(Content)
    try:
        print(f"Saving to: {outputpath}")
        DOC.save(outputpath)
        print(f"✅ Data merged into {outputpath}")
        print(f"   Exists after save? {'Yes' if outputpath.exists() else 'No'}")

    except PermissionError:
        print("❌ Could not save output.docx. Is it currently open in Word? Close it and run again.")
        raise

    except Exception:
        print("❌ Unexpected error while saving:")
        print(traceback.format_exc())
        raise            

if __name__== "__main__":
    Merge_Files(Path("Projects/File Handling Project/Source Files"), Path("Projects/File Handling Project/Output.docx"))
    print("Data merged successfully!!!!!!")    
            