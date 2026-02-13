# import os 
# from pathlib import Path 
# from PyPDF2 import PdfReader
# from openpyxl import load_workbook
# from docx import document

# extract_partial=True  # True = fetch only some lines, False = fetch full content 
# max_lines=5  # Limit for partial extraction

# def readDocx(filepath:Path) -> str:
#     DOCX = document(filepath)
#     # ContentDOCX = [p.text for p in DOCX.paragraphs if p.text.strip()]
#     ContentDOCX=[] 
#     for i in DOCX.paragraphs:
#         if i.text.strip()!="":
#             ContentDOCX.append(i.text)
#     # return "\n".join(ContentDOCX[:max_lines]) if extract_partial else"\n".join(ContentDOCX)
#     if extract_partial:
#         result=[]
#         count=0
#         for i in ContentDOCX:
#             if count<max_lines:
#                 result.append(i)
#                 count=count+1
#             else: 
#                 break
#         return"\n".join(result)
#     else:
#         return"\n".join(ContentDOCX)
# readDocx()

import os
from pathlib import Path
from PyPDF2 import PdfReader
from openpyxl import load_workbook
from docx import Document  


extract_partial = True 
max_lines = 5  



def readDocx(filepath: Path) -> str:
    if not filepath.suffix.lower() == ".docx":
        raise ValueError("Only .docx files are supported, not .doc")
    DOCX = Document(filepath)


    
    ContentDOCX = []
    for i in DOCX.paragraphs:
        if i.text.strip() != "":
            ContentDOCX.append(i.text)


    if extract_partial:
        result = []
        count = 0
        for i in ContentDOCX:
            if count < max_lines:
                result.append(i)
                count = count + 1
            else:
                break
        return "\n".join(result)
    else:
        return "\n".join(ContentDOCX)



if __name__ == "__main__": #ensure that the code inside it will not run if the file is imported elswhere
    file_path = Path("Projects/File Handling Project/Source Files/Sample.docx")  
    if file_path.exists():
        print(readDocx(file_path))
    else:
        print("File not found:", file_path)